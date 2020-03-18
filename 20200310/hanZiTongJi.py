# -*- coding: utf-8 -*-
# @Time : 2020/3/17 17:20
# @Author : 52595
# @File : hanZiTongJi.py
# @Python Version : 3.7.4
# @Software: PyCharm
"""
    功能说明：
    遍历word文档中的汉字，统计汉字出现的频次
"""
import docx
import openpyxl
from datetime import datetime


# 往Excel中写内容，参数1：excel写入的路径；参数2：excel工作表名称；参数3：需要写入的具体值
# 注意事项：如果执行代码前打开该Excel，则执行代码会报错
def write_excel_xlsx(path, sheet_name, value):
    index = len(value)
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = sheet_name
    # 对于excel格式的设置还需在研究  20200119
    '''
     # 设置列头格式
    style = workbook.XFStyle()  # 格式信息
    alignment = workbook.Alignment()  # 设置字体在单元格的位置
    alignment.vert = workbook.Alignment.VERT_CENTER  # 竖直方向
    style.alignment = alignment
    font = workbook.Font()  # 字体基本设置
    font.name = u'微软雅黑'

    sheet.write(1, 2, '汉字', style)
    sheet.write(1, 3, '频次', style)
    '''
    # 列头
    sheet.cell(row=1, column=2, value='汉字')
    sheet.cell(row=1, column=3, value='频次')
    # 具体值
    for k in range(0, len(dic)):
        sheet.cell(row=k + 2, column=2, value=dic[k][0])
        sheet.cell(row=k + 2, column=3, value=dic[k][1])
    workbook.save(path)
    print("xlsx格式表格写入数据成功！")


file = docx.Document("E:\\234.docx")
# print("段落数:"+str(len(file.paragraphs)))
# #段落数为13，每个回车隔离一段
# 输出每一段的内容
'''
for para in file.paragraphs:
    print(para.text)
    # 输出段落编号及段落内容
    for i in range(len(file.paragraphs)):
        print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)
'''
dic = {}
for i in range(len(file.paragraphs)):
    # print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
    para_single = file.paragraphs[i].text
    for j in para_single:
        if '\u4e00' <= j <= '\u9fff':
            if j in dic.keys():
                dic[j] += 1
            else:
                if j == ' ' or j == '\n':
                    continue
                else:
                    dic[j] = 1
dic = sorted(dic.items(), key=lambda d: d[1], reverse=True)
'''
for k in range(len(dic)):
    print(dic[k][0], dic[k][1])
'''
# 生成文件名
xmls_create_time = datetime.now().__format__('_%Y%m%d_%H%M_%S')
# 文件名
xmls_name = '字数使用频次统计_' + xmls_create_time + '.xlsx'
# 存放路径
xmls_path = 'D:\\developmentSoft\\pythonWorkspace\\trainInfo\\'
# 拼接后的Excle路径以及名称
name_xlsx = xmls_path + xmls_name
print('输出文件路径： ' + name_xlsx)

# 工作表的名称
sheet_name_xlsx = "汉字使用频率"
# 由于excel打开，重新运行程序，则报错，因此文件名根据时间自动生成
write_excel_xlsx(name_xlsx, sheet_name_xlsx, dic)
